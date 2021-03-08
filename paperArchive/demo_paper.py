# Copyright 2019 Lenovo(Beijing) Technologies Co.,LTD.
# Smart Education Engine Team, All Rights Reserved.
# Author: Coco Gao
# Date: 2021/1/12 下午5:01
# Email: gaojh4@lenovo.com
# Description: Demo paper split and classification.
import os
from docx import Document

import docx2txt
from paper_split_tf import PaperSplitEngine
from classify_paper_tf import ClassifyPaperEngine
import time

def file_name(path, type):
    L = []
    for root, dirs, files in os.walk(path):
        for file in files:
            if os.path.splitext(file)[1] == '.' + type:
                L.append(os.path.join(root, file))
    return L

def genertate_docx(output_text,savePath):
    # if not os.path.exists(output_file):
    #     os.makedirs(output_file)
    # my_output='\n'.join(output_text)
    # print(my_output)
    document = Document()
    # document.add_heading('Regular Labeled')
    document.add_paragraph(output_text)
    # path,file_output=os.path.split(doc_file)
    # pdb.set_trace()
    document.save(savePath)

class DemoPaper():
    def __init__(self):
        # init split engine.
        self.split_engine = PaperSplitEngine()
        # self.class_engine = ClassifyPaperEngine()

    def process_demo(self, filePath, savePath):
        lines = []
        try:
            my_text = docx2txt.process(filePath)
            my_text_line = my_text.split('\n')

            for line in my_text_line:
                # if ('试卷答案' in line) or ('参考答案' in line):
                #     continue
                line_no_blank = line.strip()
                if len(line_no_blank) > 0:
                    # print(line_no_blank)
                    lines.append(line_no_blank)

            predict_result = self.split_engine.fast_demo(lines)
            for i,pr in enumerate(predict_result):
                print(lines[i], pr)
            # # class
            title_list = []
            num_one = 0
            title = ""
            for i,line in enumerate(lines):
                if predict_result[i] == "1":
                    if num_one != 0:
                        title_list.append(title)
                        title = line
                    else:
                        title = title + "\n" + line
                    num_one = num_one + 1
                else:
                    title = title + "\n" + line
                if i == len(lines) - 1:
                    title_list.append(title)
            an_list = self.class_engine.fast_demo(title_list)
            print(an_list)
            word_txt = ""
            for i, an in enumerate(an_list):
                print(title_list[i], an)
                word_txt += title_list[i]
                word_txt += "---------------------------切分线，标注：-------------------------------------"
                word_txt += "\t"
                word_txt += an
                word_txt += "\n"
                word_txt += "---------------------------切分线，标注：------------------------------------"
                word_txt += "\n"
                word_txt += "\n"
                word_txt += "\n"
                # print("\n")
                # print("\n")
            genertate_docx(word_txt, savePath)
            # print(word_txt)
        except:
            print("file is error...")
            return "exception..."


    def process_split(self, filePath):
        files = file_name(filePath, "docx")
        max_acc = [0]
        max_f = ""
        rlt = []

        for f in files:
            print(f)
            try:
                my_text = docx2txt.process(f)
                my_text_line = my_text.split('\n')
                lines = []

                for line in my_text_line:
                    if ('试卷答案' in line) or ('参考答案' in line):
                        break
                    line_no_blank = line.strip()
                    if len(line_no_blank) > 0:
                        # print(line_no_blank)
                        lines.append(line_no_blank)

                acc = self.split_engine.model_eval_word(lines)
                rlt.append([f, acc])

                if acc > max_acc:
                    max_acc = acc
                    max_f = f
            except:
                print(f, "is exception file")

        for e in rlt:
            print(e)
        print(max_f, max_acc)

    def eval_docs(self, filePath):
        files = file_name(filePath, "docx")
        lines = []

        for f in files:
            print(f)
            try:
                my_text = docx2txt.process(f)
                my_text_line = my_text.split('\n')

                for line in my_text_line:
                    if ('试卷答案' in line) or ('参考答案' in line):
                        break
                    line_no_blank = line.strip()
                    if len(line_no_blank) > 0:
                        # print(line_no_blank)
                        lines.append(line_no_blank)
            except:
                print(f, "is exception file")

        acc = self.split_engine.model_eval_word(lines)
        print(len(lines), acc)


if __name__ == '__main__':

    # filePath = ""

    s1 = time.time()
    demo = DemoPaper()
    s2 = time.time()

    print("init time:", str(s2-s1))

    # src_path = "/home/jhgao/Desktop/paperTest/"
    # demo.process_split(src_path)
    # src_path = "/home/jhgao/Desktop/paperTest/"
    # demo.process_demo("7c4b92993f2d4ffbad6c031a9d37a942.docx","/home/jhgao/Desktop/切分后_测试集_初中英语.docx")
    # demo.process_demo("/home/jhgao/Desktop/网络试卷_初中英语.docx", "/home/jhgao/Desktop/切分后_网络试卷_初中英语.docx")
    src_path = "/home/jhgao/Desktop/docx_paper/"
    dest_path = "/home/jhgao/Desktop/output_seg/"

    # file_list = file_name(src_path, "docx")
    # s3 = time.time()
    # for f in file_list:
    #     f_name =  f.split("/")[-1]
    #     after_name = "切分后_" + f_name
    #     print(f)
    #     demo.process_demo(src_path + f_name, dest_path + after_name)
    #
    # s4 = time.time()
    # print("average time:", (s4-s3)/len(file_list), "s")

    src_path = "/home/jhgao/Desktop/data_excel/"
    f_list = file_name(src_path, "xlsx")
    rlt = []
    for f in f_list:
        print(f)
        acc, recall = demo.split_engine.model_eval_excel(f)
        rlt.append([f, acc, recall])

    print(rlt)